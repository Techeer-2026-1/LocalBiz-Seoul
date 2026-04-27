"""OpenSearch 인덱스 + 검색 파이프라인 + 4-Tier 통합 가이드 docx 생성.

Usage:
    cd backend && source venv/bin/activate
    python scripts/generate_os_guide.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    shading.append(elm)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "2B579A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    return table


def build():
    doc = Document()
    REPO = Path(__file__).resolve().parents[2]

    # ─── Title ───
    t = doc.add_heading("LocalBiz Intelligence", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_heading("OpenSearch 인덱스 + 검색 파이프라인 + 데이터 계층\n통합 가이드", level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("작성일: 2026-04-13 | v1.0 | 작성: 이정 (BE/PM) + Claude").font.size = Pt(9)

    # ─── 목차 ───
    doc.add_page_break()
    doc.add_heading("목차", level=1)
    toc = [
        "1. 개요 — PostgreSQL과 OpenSearch의 역할 분담",
        "2. PG ↔ OS 동기화 아키텍처",
        "3. OpenSearch 인덱스 상세 (3개)",
        "  3.1 places_vector — 장소 의미 검색",
        "  3.2 events_vector — 행사 의미 검색",
        "  3.3 place_reviews — 리뷰 키워드 검색",
        "4. 검색 파이프라인 (6단계)",
        "  4.1 Intent Router",
        "  4.2 SQL 필터",
        "  4.3 Vector k-NN",
        "  4.4 BM25 보조",
        "  4.5 place_reviews 보강",
        "  4.6 LLM Rerank",
        "5. 4-Tier 데이터 계층",
        "6. page_content 3-Layer 구조",
        "7. 임베딩 모델 + 분석기 설정",
        "8. 기능별 검색 경로 매핑",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Bullet")

    # ─── 1. 개요 ───
    doc.add_page_break()
    doc.add_heading("1. 개요 — PostgreSQL과 OpenSearch의 역할 분담", level=1)

    doc.add_paragraph(
        "LocalBiz Intelligence는 하이브리드 검색 아키텍처를 사용한다. "
        "PostgreSQL(Cloud SQL)은 정형 데이터 저장 + SQL 필터링을, "
        "OpenSearch(GCE)는 벡터 의미 검색 + 키워드 매칭을 담당한다."
    )

    add_table(
        doc,
        ["역할", "PostgreSQL", "OpenSearch"],
        [
            ["저장", "정형 테이블 12개 (535K places 등)", "벡터 인덱스 3개 (768d 임베딩)"],
            ["검색 방식", "SQL WHERE + PostGIS 공간 쿼리", "k-NN 코사인 유사도 + BM25 키워드"],
            ["강점", "정확한 필터링 (category, district, 날짜)", "의미적 유사도 ('분위기 좋은', '카공하기 좋은')"],
            ["약점", "'분위기 좋은' 같은 비정형 쿼리 불가", "정확한 범위 필터 어려움"],
            ["연결", "place_id / event_id", "_id (동일 값으로 application-level 연결)"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "핵심 원칙: 두 DB는 place_id = _id (또는 event_id = _id)로 연결된다. "
        "검색 시 양쪽에서 후보를 가져온 뒤 합산·중복 제거하여 최종 결과를 만든다."
    )

    # ─── 2. 동기화 ───
    doc.add_page_break()
    doc.add_heading("2. PG ↔ OS 동기화 아키텍처", level=1)

    # 다이어그램 이미지 삽입
    sync_img = REPO / "기획" / "OS_PG_동기화_다이어그램.png"
    if sync_img.exists():
        doc.add_picture(str(sync_img), width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("동기화 규칙", level=2)

    rules = [
        (
            "places → places_vector",
            "places 테이블의 각 row를 page_content(자연어 텍스트)로 변환 → "
            "Gemini gemini-embedding-001로 768d 벡터 생성 → OpenSearch에 _id=place_id로 적재. "
            "places 변경 시 places_vector도 재적재 필요.",
        ),
        (
            "events → events_vector",
            "events 테이블의 title + summary를 결합 → Gemini 768d 임베딩 → OpenSearch에 _id=event_id로 적재.",
        ),
        (
            "places → place_reviews",
            "Naver Blog Search API로 장소별 리뷰 크롤링 → Gemini로 6지표 채점 + 요약 → "
            "summary_text 임베딩 → OpenSearch에 _id=review_{place_id}로 적재. "
            "이 인덱스는 places와 직접 동기화가 아니라 별도 크롤링 파이프라인으로 채워짐.",
        ),
    ]
    for title_text, desc in rules:
        p = doc.add_paragraph()
        p.add_run(title_text + ": ").bold = True
        p.add_run(desc)

    # ─── 3. 인덱스 상세 ───
    doc.add_page_break()
    doc.add_heading("3. OpenSearch 인덱스 상세 (3개)", level=1)

    doc.add_paragraph(
        "공통 설정: 768d (Gemini gemini-embedding-001), nori 한국어 분석기, k-NN HNSW cosinesimil, shard 1, replica 0"
    )

    # 3.1 places_vector
    doc.add_heading("3.1 places_vector — 장소 의미 검색", level=2)
    doc.add_paragraph("535,431 docs. PG places 테이블과 1:1 매핑.")

    add_table(
        doc,
        ["필드", "타입", "용도", "값 예시"],
        [
            ["_id", "(document ID)", "PG places.place_id와 동일", "MA010120220804265295"],
            ["place_id", "keyword", "PG 연결키 (필터용)", "MA010120220804265295"],
            ["name", "text (nori)", "장소명 키워드 검색", "60계치킨암사"],
            [
                "page_content",
                "text (nori)",
                "3-Layer 자연어 텍스트. 벡터 임베딩의 원본이자 BM25 키워드 검색 대상",
                "강동구 암사2동에 위치한 치킨 전문점. 60계치킨암사. 업종: 치킨. 혼밥이나 데이트에 적합...",
            ],
            [
                "embedding",
                "knn_vector 768d",
                "page_content의 Gemini 임베딩 벡터. k-NN 코사인 유사도 검색에 사용",
                "[0.012, -0.034, 0.087, ...]",
            ],
            ["category", "keyword", "SQL 필터 대응. Intent Router가 추론한 카테고리로 pre-filtering", "음식점"],
            ["sub_category", "keyword", "세부 분류 필터", "기타 간이"],
            ["district", "keyword", "자치구 필터", "강동구"],
            ["source", "keyword", "데이터 출처 필터", "sosang_biz_202512"],
            ["lat", "float", "위도 (지도 마커 표시용)", "37.5508"],
            ["lng", "float", "경도", "127.1269"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("page_content가 핵심. ").bold = True
    p.add_run(
        "이 필드의 품질이 벡터 검색 정확도를 결정한다. "
        "현재 3-Layer 구조: Layer 1(구조적 속성) + Layer 2(카테고리 기본 설명 60종) + Layer 3(리뷰 요약). "
        "상세는 §6 참조."
    )

    # 3.2 events_vector
    doc.add_heading("3.2 events_vector — 행사 의미 검색", level=2)
    doc.add_paragraph("7,301 docs. PG events 테이블과 1:1 매핑.")

    add_table(
        doc,
        ["필드", "타입", "용도", "값 예시"],
        [
            ["_id", "(document ID)", "PG events.event_id와 동일", "evt-abc123"],
            ["event_id", "keyword", "PG 연결키", "evt-abc123"],
            ["title", "text (nori)", "행사명 키워드 검색", "2026 핸드아티코리아"],
            [
                "description",
                "text (nori)",
                "title + summary 결합 텍스트. 벡터 원본 + BM25 대상",
                "2026 핸드아티코리아. 코엑스전시장 B홀 전시/미술",
            ],
            ["embedding", "knn_vector 768d", "description의 Gemini 임베딩", "[0.023, ...]"],
            ["category", "keyword", "행사 분류 필터", "전시/미술"],
            ["district", "keyword", "자치구 필터", "강남구"],
            ["date_start", "date", "날짜 범위 pre-filtering. '이번 주말 전시' 쿼리 시 range filter 적용", "2026-04-15"],
            ["date_end", "date", "종료일 필터", "2026-04-20"],
            ["source", "keyword", "데이터 출처", "서울시문화행사"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("date_start/date_end 필터가 핵심 차별점. ").bold = True
    p.add_run("벡터 검색 전에 날짜 범위로 pre-filtering하여 관련 없는 지난 행사를 제외한다.")

    # 3.3 place_reviews
    doc.add_heading("3.3 place_reviews — 리뷰 키워드 검색", level=2)
    doc.add_paragraph("~7K docs (적재 진행 중). PG places와 app-level 연결 (FK 아님).")

    add_table(
        doc,
        ["필드", "타입", "용도", "값 예시"],
        [
            ["_id", "(document ID)", "review_{place_id} 형식", "review_MA010120220804265295"],
            ["review_id", "keyword", "리뷰 고유 ID", "review_MA010120220804265295"],
            ["place_id", "keyword", "PG places 연결키", "MA010120220804265295"],
            ["place_name", "text (nori)", "장소명 (검색 편의)", "60계치킨암사"],
            [
                "summary_text",
                "text (nori)",
                "리뷰 요약 + 키워드 결합 텍스트. 벡터 원본",
                "바삭하고 촉촉한 식감과 적당한 간으로 맥주와 완벽한 조합 혼밥 맥주안주 바삭촉촉",
            ],
            ["embedding", "knn_vector 768d", "summary_text의 Gemini 임베딩", "[0.045, ...]"],
            [
                "keywords",
                "keyword (array)",
                "경험적 키워드 배열. 태그 필터링 가능",
                '["혼밥", "맥주안주", "바삭촉촉", "가성비"]',
            ],
            ["stars", "float", "6지표 평균 점수 (1~5)", "4.2"],
            ["source", "keyword", "수집 소스", "naver_blog_batch"],
            ["category", "keyword", "장소 카테고리 (PG에서 복사)", "음식점"],
            ["district", "keyword", "자치구", "강동구"],
            ["analyzed_at", "date", "분석 일시", "2026-04-13"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("keywords 배열이 핵심. ").bold = True
    p.add_run(
        "'혼밥하기 좋은 식당' 쿼리 시 keywords에 '혼밥'이 포함된 place_reviews를 찾고, "
        "해당 place_id로 PG places를 JOIN하여 상세 정보를 반환한다. "
        "이 인덱스는 places_vector의 page_content만으로 부족한 '경험적 쿼리'를 보강한다."
    )

    # ─── 4. 검색 파이프라인 ───
    doc.add_page_break()
    doc.add_heading("4. 검색 파이프라인 (6단계)", level=1)

    # 다이어그램 이미지
    search_img = REPO / "기획" / "검색_파이프라인_다이어그램.png"
    if search_img.exists():
        doc.add_picture(str(search_img), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("사용자 쿼리 '카공하기 좋은 카페 성수동'을 예시로 각 단계를 설명한다.")

    # 4.1
    doc.add_heading("4.1 Intent Router (Gemini Flash)", level=2)
    doc.add_paragraph(
        "사용자 메시지를 Gemini 2.5 Flash에 전달하여 13개 intent 중 하나를 분류한다. "
        "동시에 카테고리를 추론('카페')하고, 쿼리를 확장한다 "
        "('카공' → '콘센트 카페', '스터디 카페', '넓은 테이블'). "
        "이 확장된 쿼리가 이후 단계의 입력이 된다."
    )
    add_table(
        doc,
        ["입력", "출력", "소요"],
        [
            [
                "'카공하기 좋은 카페 성수동'",
                "intent=PLACE_RECOMMEND, category='카페', district='성동구', "
                "expanded=['콘센트 카페', '스터디 카페', '넓은 테이블']",
                "~300ms",
            ],
        ],
    )

    # 4.2
    doc.add_heading("4.2 SQL 필터 (PG + PostGIS)", level=2)
    doc.add_paragraph(
        "Intent Router가 추론한 category와 district로 PG places 테이블을 SQL 필터링한다. "
        "PostGIS ST_DWithin으로 반경 검색도 가능. 정확한 구조적 매칭 담당."
    )
    doc.add_paragraph(
        "예: SELECT * FROM places WHERE category='카페' AND district LIKE '%성동%' "
        "AND ST_DWithin(geog, ST_MakePoint(127.06, 37.54)::geography, 2000) LIMIT 50"
    )

    # 4.3
    doc.add_heading("4.3 Vector k-NN (OpenSearch places_vector)", level=2)
    doc.add_paragraph(
        "확장된 쿼리('콘센트 카페 스터디 카페 넓은 테이블')를 Gemini 768d로 임베딩한 뒤, "
        "OpenSearch places_vector에서 k-NN 코사인 유사도 검색을 수행한다. "
        "category='카페' 필터를 함께 적용하여 카테고리 내에서만 의미 매칭한다."
    )
    doc.add_paragraph(
        "page_content에 '카공하기 좋은 조용한 분위기'가 포함된 카페가 높은 유사도를 얻는다. "
        "이것이 Layer 2 카테고리 기본 설명의 가치 — 모든 카페에 '카공' 키워드가 포함되어 있으므로 "
        "최소한의 매칭이 보장된다."
    )

    # 4.4
    doc.add_heading("4.4 BM25 보조 (OpenSearch nori)", level=2)
    doc.add_paragraph(
        "벡터 검색과 함께 page_content의 키워드 매칭(BM25)도 수행한다. "
        "'성수동'이라는 지역명이 page_content 주소에 포함된 문서가 BM25에서 높은 점수를 얻는다. "
        "벡터 검색이 의미적 유사도를, BM25가 키워드 정확도를 보완하는 하이브리드 구조."
    )

    # 4.5
    doc.add_heading("4.5 place_reviews 보강", level=2)
    doc.add_paragraph(
        "place_reviews 인덱스에서 keywords 배열에 '카공'이 포함된 리뷰를 검색한다. "
        "해당 장소의 실제 리뷰에서 '콘센트 많고 넓어서 작업하기 좋아요' 같은 경험적 정보가 있으면 "
        "해당 place_id에 가산점을 부여한다. "
        "place_reviews가 없는 장소는 이 단계에서 불이익 없음 (Layer 1+2만으로도 기본 매칭)."
    )

    # 4.6
    doc.add_heading("4.6 LLM Rerank (Gemini Flash)", level=2)
    doc.add_paragraph(
        "위 4단계에서 모은 후보 ~50건의 page_content를 Gemini Flash에 전달하여 "
        "'카공하기 좋은 카페 성수동'에 가장 적합한 순서로 재정렬한다. "
        "LLM은 상식을 활용하여 '콘센트', '넓은 좌석', '조용한 분위기'를 이해하고 판단한다. "
        "이 단계가 벡터 검색만으로 해결하기 어려운 '야경이 예쁜 곳' 같은 추상적 쿼리를 커버한다."
    )
    add_table(
        doc,
        ["입력", "출력", "소요"],
        [["후보 50건 page_content + 원본 쿼리", "Top 5 순위 재정렬", "~1-2초"]],
    )

    # ─── 5. 4-Tier ───
    doc.add_page_break()
    doc.add_heading("5. 4-Tier 데이터 계층", level=1)

    tier_img = REPO / "기획" / "4Tier_데이터_계층_다이어그램.png"
    if tier_img.exists():
        doc.add_picture(str(tier_img), width=Inches(3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tiers = [
        (
            "Tier 1: 구조적 (535K, 100% 커버)",
            "CSV raw_data 속성 + Layer 2 카테고리 기본 설명",
            "'강남 치킨집', '무료 주차장', '24시간 약국'",
            "모든 장소가 기본적인 벡터 매칭 가능. 카테고리+지역+상호명+업종 정보.",
        ),
        (
            "Tier 2: 경험적 (~7K)",
            "Naver Blog 리뷰 배치 크롤링 → Gemini 6지표 + 키워드",
            "'분위기 좋은 카페', '사장님이 친절한 식당'",
            "리뷰 기반 경험적 키워드. place_reviews 인덱스에 저장. "
            "리뷰가 있는 장소만 커버되므로 Tier 1이 fallback 역할.",
        ),
        (
            "Tier 3: LLM Rerank (런타임)",
            "Gemini Flash 상식 기반 후보 재정렬",
            "'야경이 예쁜 곳', '혼자 여행하기 좋은 서울 관광지'",
            "벡터 검색으로 후보를 모은 뒤, LLM이 상식으로 판단하여 순위 조정. 카테고리 추론이 어려운 추상적 쿼리 대응.",
        ),
        (
            "Tier 4: 큐레이션 (~1K)",
            "YouTube 전문 + transcript 구조화 추출 + wiki 엔티티",
            "'콘센트+조용+넓은 카페 성수', '혼밥+LP+아날로그 감성 카페 종로'",
            "사람이 검증한 또는 YouTube 크리에이터가 상세 소개한 장소. "
            "features controlled vocabulary (outlet, quiet 등)로 정밀 매칭. "
            "data/wiki/ 엔티티 + data/extracted/youtube/ JSON으로 관리.",
        ),
    ]

    for title_text, data, example, desc in tiers:
        doc.add_heading(title_text, level=2)
        p1 = doc.add_paragraph()
        p1.add_run("데이터: ").bold = True
        p1.add_run(data)
        p2 = doc.add_paragraph()
        p2.add_run("커버하는 쿼리: ").bold = True
        p2.add_run(example)
        doc.add_paragraph(desc)

    # ─── 6. page_content ───
    doc.add_page_break()
    doc.add_heading("6. page_content 3-Layer 구조", level=1)

    doc.add_paragraph(
        "places_vector의 page_content 필드는 벡터 검색 품질의 핵심이다. "
        "3개 Layer로 구성되며, 모든 장소가 Layer 1+2를 보유하고, "
        "리뷰가 있는 장소만 Layer 3가 추가된다."
    )

    add_table(
        doc,
        ["Layer", "내용", "커버리지", "생성 방법", "역할"],
        [
            [
                "Layer 1: 구조적",
                "상호명, 카테고리, 소분류, 지역, 주소, raw_data 속성 (면적/좌석수/진료시간 등)",
                "535K (100%)",
                "source별 특화 템플릿 20종 (page_content.py)",
                "'강남 치킨집' 구조적 매칭",
            ],
            [
                "Layer 2: 카테고리 기본 설명",
                "'혼밥이나 데이트에 적합. 카공하기 좋은 조용한 분위기...'",
                "535K (100%)",
                "Gemini 1회 생성 → 60종 JSON (category_descriptions.json)",
                "'카공 카페' 같은 경험적 쿼리에도 최소 매칭 보장",
            ],
            [
                "Layer 3: 리뷰 요약",
                "'바삭한 치킨, 맥주와 완벽한 조합. 주차 가능.'",
                "~7K (크롤링 성공분)",
                "Naver Blog 배치 크롤링 → Gemini 요약",
                "실제 경험 기반 정밀 매칭",
            ],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("page_content 실제 예시", level=2)

    examples = [
        (
            "소상공인 음식점 (Layer 1+2)",
            "강동구 암사2동에 위치한 치킨 전문점. 60계치킨암사. 업종: 치킨. 분류: 치킨 전문점. "
            "지점: 선사점. 암사동정웅빌딩 1층. "
            "전통적인 한식의 깊은 맛과 정갈함을 느낄 수 있는 곳입니다. "
            "든든한 혼밥이나 소중한 사람과의 데이트에 아늑하고 특별한 분위기를 더해줍니다. "
            "서울특별시 강동구 상암로3길 8.",
        ),
        (
            "병의원 (Layer 1, raw_data 풍부)",
            "금천구에 위치한 의원 의료. 가산기대찬의원. 병원분류: 의원. "
            "삼성서울병원 외래교수 출신 구강외과 전문의 진료, 진료과목 - 임플란트, "
            "치조골 뼈이식 수술, 매복 사랑니 발치, 턱관절 악관절 질환의 치료. "
            "진료시간: 월 09:00-19:30, 화 09:00-19:30... 응급실 미운영.",
        ),
        (
            "공영주차장 (Layer 1, 요금 정보)",
            "구로구에 위치한 노외 주차장. 구로디지털단지역 공영주차장. "
            "유료. 총 91면. 기본 5분 320원. 평일 00:00-24:00. 토요일 무료. 공휴일 무료.",
        ),
    ]

    for title_text, content in examples:
        p = doc.add_paragraph()
        p.add_run(title_text + "\n").bold = True
        p.add_run(content).font.size = Pt(8)
        doc.add_paragraph()

    # ─── 7. 임베딩 + 분석기 ───
    doc.add_page_break()
    doc.add_heading("7. 임베딩 모델 + 분석기 설정", level=1)

    doc.add_heading("임베딩 모델", level=2)
    add_table(
        doc,
        ["항목", "값"],
        [
            ["모델", "Gemini gemini-embedding-001"],
            ["차원", "768d"],
            ["비용", "무료 (1,500 RPM)"],
            ["입력 제한", "2,048 토큰 (page_content 2,000자 제한)"],
            ["호출 방식", "batchEmbedContents (100건/call)"],
            ["불변식", "#7: 768d Gemini만. OpenAI 사용 시 PR 차단"],
        ],
    )

    doc.add_heading("nori 한국어 분석기", level=2)
    doc.add_paragraph(
        "OpenSearch 3 인덱스 모두 nori_analyzer를 사용한다. "
        "nori_tokenizer(한국어 형태소 분석) + lowercase 필터. "
        "'카공하기 좋은 카페'를 '카공', '좋은', '카페'로 토큰화하여 BM25 키워드 검색에 활용."
    )

    doc.add_heading("HNSW k-NN 설정", level=2)
    add_table(
        doc,
        ["항목", "값", "설명"],
        [
            ["algorithm", "hnsw", "Hierarchical Navigable Small World"],
            ["engine", "nmslib", "OpenSearch 기본 k-NN 엔진"],
            ["space_type", "cosinesimil", "코사인 유사도 (방향 기반, 크기 무관)"],
            ["ef_search", "100", "검색 시 탐색 범위 (정확도↑ 속도↓)"],
        ],
    )

    # ─── 8. 기능별 경로 ───
    doc.add_page_break()
    doc.add_heading("8. 기능별 검색 경로 매핑", level=1)

    doc.add_paragraph("각 기능(intent)이 어떤 검색 경로를 사용하는지 정리.")

    add_table(
        doc,
        ["기능 (Intent)", "SQL", "places_vector", "events_vector", "place_reviews", "LLM Rerank", "외부 API"],
        [
            ["PLACE_SEARCH", "O", "O", "", "", "O", ""],
            ["PLACE_RECOMMEND", "O", "O", "", "O", "O", "Naver Blog"],
            ["EVENT_SEARCH", "O", "", "O", "", "", "Naver fallback"],
            ["COURSE_PLAN", "O (PostGIS)", "O", "", "", "O", "OSRM"],
            ["BOOKING", "O", "", "", "", "", "Google Places"],
            ["CALENDAR", "", "", "", "", "", "Google Calendar"],
            ["REVIEW_COMPARE", "", "", "", "O", "", "Naver+Google 런타임"],
            ["COST_ESTIMATE", "O", "", "", "", "O", "Naver Blog"],
            ["CROWDEDNESS", "O (pop_stats)", "", "", "", "", "서울시 API"],
            ["IMAGE_SEARCH", "", "O", "", "", "", "Gemini Vision"],
            ["ANALYSIS", "", "", "", "O", "", "Naver+Google 런타임"],
            ["GENERAL", "", "", "", "", "", "Gemini LLM 직접"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "O = 해당 데이터 소스 사용. 빈 칸 = 미사용. "
        "대부분의 장소 관련 기능은 SQL + places_vector + LLM Rerank 3단계를 거친다. "
        "행사 검색만 events_vector를 사용하고, 리뷰 비교/분석은 place_reviews를 사용한다."
    )

    # Save
    output = REPO / "기획" / "OpenSearch_검색_파이프라인_통합가이드.docx"
    doc.save(str(output))
    print(f"saved: {output}")


if __name__ == "__main__":
    build()
