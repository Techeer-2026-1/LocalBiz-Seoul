"""ERD v6.3 최종 확정 보고서 docx 생성.

Usage:
    cd backend && source venv/bin/activate
    python scripts/generate_erd_report.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    shading.append(elm)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
        set_cell_shading(cell, "2B579A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    return table


def build():
    doc = Document()

    # ─── Title ───
    t = doc.add_heading("LocalBiz Intelligence", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_heading("ERD v6.3 최종 확정 보고서", level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "작성일: 2026-04-14 | 작성: 이정 (BE/PM) + Claude\n"
        "Cloud SQL 실측 기준 | 12 테이블 · 12 FK · 기능 47개 · API 19개 전수 검증 완료"
    ).font.size = Pt(9)

    # ─── 목차 ───
    doc.add_page_break()
    doc.add_heading("목차", level=1)
    toc = [
        "1. 개요",
        "  1.1 ERD 버전 이력",
        "  1.2 설계 원칙",
        "  1.3 PK 전략",
        "  1.4 created_at / updated_at / is_deleted 적용 기준",
        "2. 테이블 상세 (12개)",
        "  2.1 장소 (places)",
        "  2.2 행사 (events)",
        "  2.3 행정동 (administrative_districts)",
        "  2.4 생활인구 (population_stats)",
        "  2.5 행정동코드매핑 (admin_code_aliases)",
        "  2.6 사용자 (users)",
        "  2.7 OAuth토큰 (user_oauth_tokens)",
        "  2.8 대화 (conversations)",
        "  2.9 메시지 (messages)",
        "  2.10 북마크 (bookmarks)",
        "  2.11 공유링크 (shared_links)",
        "  2.12 피드백 (feedback)",
        "3. FK 관계 상세 (12개)",
        "4. OpenSearch 인덱스 (3개)",
        "5. 정규화 판단",
        "6. 기능·API 매핑 검증",
        "7. 부록: LangGraph 자동 테이블",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Bullet")

    # ─── 1. 개요 ───
    doc.add_page_break()
    doc.add_heading("1. 개요", level=1)

    doc.add_heading("1.1 ERD 버전 이력", level=2)
    add_table(
        doc,
        ["버전", "날짜", "주요 변경"],
        [
            ["v6.0", "2026-04-09", "초기 12 테이블 설계 (place_analysis 포함)"],
            ["v6.1", "2026-04-11", "ERD 상세설명보고서 docx 작성"],
            ["v6.2", "2026-04-12", "place_analysis 컬럼 정합, bookmarks/shared_links/feedback DDL"],
            [
                "v6.3 → v6.1 확정",
                "2026-04-13",
                "place_analysis DROP, admin_code_aliases/user_oauth_tokens 추가, "
                "geom/geog PostGIS 전환, 복합키→BIGSERIAL, is_deleted BOOLEAN 통일, "
                "shared_links FK 추가, JSONB 표기 통일",
            ],
            ["v6.3", "2026-04-14", "v6.1 확정본 갱신: 제목/FK/비정규화 v6.3 반영, docx 재생성"],
        ],
    )

    doc.add_heading("1.2 설계 원칙", level=2)
    principles = [
        '"공통 기둥 + 유연한 주머니(JSONB)" 패턴 — 160종 CSV 스키마를 raw_data JSONB로 통합',
        "PostgreSQL place_id = OpenSearch _id — application 레벨에서 동일 키로 두 DB 연결",
        "대화 이력 이원화 — LangGraph checkpoint(LLM 컨텍스트, 압축 가능) + messages(UI 원본, append-only)",
        "3NF 준수 + 성능 위한 의도적 비정규화 3건 허용 (v6.3: place_analysis DROP으로 4→3)",
        "append-only 테이블 — messages, population_stats, feedback (UPDATE/DELETE 금지)",
    ]
    for p in principles:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("1.3 PK 전략", level=2)
    add_table(
        doc,
        ["테이블", "PK 타입", "PK 컬럼", "이유"],
        [
            ["places", "VARCHAR(36)", "place_id", "OpenSearch places_vector._id와 동일 (소스별 자연키 or UUID)"],
            ["events", "VARCHAR(36)", "event_id", "OpenSearch events_vector._id와 동일"],
            ["administrative_districts", "VARCHAR(20) 자연키", "adm_dong_code", "행정동 코드 자체가 유일"],
            ["population_stats ~ feedback", "BIGSERIAL", "id 등", "내부 테이블, JOIN 효율"],
            ["admin_code_aliases", "BIGSERIAL", "id", "UNIQUE(old_code, new_code) 별도"],
        ],
    )

    doc.add_heading("1.4 created_at / updated_at / is_deleted 적용 기준", level=2)
    add_table(
        doc,
        ["테이블", "created_at", "updated_at", "is_deleted", "이유"],
        [
            ["places", "O", "O", "O (BOOLEAN)", "소프트 삭제 지원"],
            ["events", "O", "O", "O", "소프트 삭제"],
            ["administrative_districts", "O", "O", "X", "마스터 데이터, 삭제 불허 (RESTRICT)"],
            ["population_stats", "O", "X", "X", "append-only 시계열"],
            ["admin_code_aliases", "O", "X", "X", "참조 테이블"],
            ["users", "O", "O", "O", "계정 비활성화"],
            ["user_oauth_tokens", "O", "O", "O", "토큰 폐기 이력"],
            ["conversations", "O", "O", "O", "대화 삭제"],
            ["messages", "O", "X", "X", "append-only (절대 삭제/수정 불가)"],
            ["bookmarks", "O", "O", "O", "북마크 삭제"],
            ["shared_links", "O", "O", "O", "공유 해제"],
            ["feedback", "O", "X", "X", "append-only (피드백 이력)"],
        ],
    )

    # ─── 2. 테이블 상세 ───
    doc.add_page_break()
    doc.add_heading("2. 테이블 상세 (12개)", level=1)

    TABLES = [
        {
            "num": "2.1",
            "name_ko": "장소",
            "name_en": "places",
            "rows": "535,431",
            "desc": "서울시 공공 CSV 153종 + 소상공인 상가(상권) + 서울 열린데이터 API + Google 지오코딩을 통합한 장소 마스터 테이블. "
            "18 카테고리(v0.2 분류표), 48 source, 25 서울 자치구를 커버한다.",
            "features": "장소 검색(PLACE_SEARCH), 장소 추천(PLACE_RECOMMEND), 코스 추천(COURSE_PLAN), 예약 연동(BOOKING), 지도 마커, 비용 견적(COST_ESTIMATE)",
            "how": "사용자가 '강남 카페 추천'이라고 하면 Intent Router가 category='카페'를 추론하고, "
            "SQL WHERE category='카페' AND district LIKE '%강남%' + PostGIS ST_DWithin 반경 검색으로 후보를 추출한다. "
            "동시에 OpenSearch places_vector에서 '카페 추천' 벡터 k-NN 검색으로 의미 매칭 후보를 가져온다. "
            "두 결과를 합산·중복 제거 후 LLM Rerank(Gemini Flash)로 최종 순위를 정하여 places[] 블록으로 응답한다.",
            "columns": [
                [
                    "place_id",
                    "장소ID",
                    "VARCHAR(36)",
                    "PK",
                    "",
                    "NO",
                    "",
                    "UUID 또는 소스별 자연키 (소상공인: 상가업소번호 MA01...)",
                ],
                ["name", "상호명", "VARCHAR(200)", "", "", "NO", "", "상호명/시설명. 검색 키워드 매칭 대상"],
                [
                    "category",
                    "대분류",
                    "VARCHAR(50)",
                    "",
                    "",
                    "NO",
                    "",
                    "v0.2 18종 고정 enum. validate_category() 함수로 강제",
                ],
                [
                    "sub_category",
                    "소분류",
                    "VARCHAR(100)",
                    "",
                    "",
                    "YES",
                    "",
                    "source별 세부 분류. SQL 필터 + 벡터 검색 page_content에 포함",
                ],
                ["address", "주소", "TEXT", "", "", "YES", "", "도로명 또는 지번"],
                [
                    "district",
                    "자치구",
                    "VARCHAR(50)",
                    "",
                    "",
                    "NO",
                    "",
                    "비정규화. 25 서울 자치구. 매번 공간 JOIN 대신 직접 필터",
                ],
                [
                    "geom",
                    "좌표",
                    "geometry(Point,4326)",
                    "",
                    "",
                    "YES",
                    "",
                    "PostGIS WGS84. 평면 기하 연산용 (ST_Contains 행정동 매칭)",
                ],
                ["phone", "전화번호", "VARCHAR(20)", "", "", "YES", "", "연락처. 장소 상세 UI에 표시"],
                [
                    "google_place_id",
                    "구글장소ID",
                    "VARCHAR(100)",
                    "",
                    "",
                    "YES",
                    "",
                    "deprecated. 향후 Google Places API 캐시 가능성으로 유지",
                ],
                [
                    "booking_url",
                    "예약URL",
                    "TEXT",
                    "",
                    "",
                    "YES",
                    "",
                    "사전 확보 URL 캐시. NULL이면 런타임 Google/Naver 딥링크 생성",
                ],
                [
                    "raw_data",
                    "원본데이터",
                    "JSONB",
                    "",
                    "",
                    "YES",
                    "",
                    "CSV 전체 row 보존. 160종 CSV 스키마 통합 (source별 상이)",
                ],
                [
                    "source",
                    "데이터출처",
                    "VARCHAR(50)",
                    "",
                    "",
                    "NO",
                    "",
                    "48종 source 태그 (sosang_biz_202512, seoul_park 등)",
                ],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "YES", "CURRENT_TIMESTAMP", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "YES", "CURRENT_TIMESTAMP", "ETL 재적재 시 갱신"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "소프트 삭제"],
                [
                    "geog",
                    "지리좌표",
                    "geography",
                    "",
                    "",
                    "YES",
                    "GENERATED",
                    "geom에서 자동 생성. 구면 거리 연산용 (ST_DWithin 미터 단위 정확 거리)",
                ],
            ],
        },
        {
            "num": "2.2",
            "name_ko": "행사",
            "name_en": "events",
            "rows": "7,301",
            "desc": "서울시 문화행사, 공공시설 예약, 구별 축제 등을 통합한 행사/축제 테이블. 8 source.",
            "features": "행사 검색(EVENT_SEARCH) — 정형(DB 우선→Naver fallback) + 비정형(events_vector k-NN), 코스 추천(COURSE_PLAN), 지도 마커",
            "how": "'이번 주말 서울 전시회' 쿼리 시 date_start/date_end 범위 필터(SQL) + events_vector 의미 검색(OS) 병합. "
            "결과 부족 시 Naver 뉴스+블로그 병렬 6req → Gemini JSON mode 구조화 추출 → confidence 필터로 보충.",
            "columns": [
                ["event_id", "행사ID", "VARCHAR(36)", "PK", "", "NO", "gen_random_uuid()", "UUID"],
                ["title", "행사명", "VARCHAR(200)", "", "", "NO", "", "행사 제목. 벡터 임베딩 대상"],
                ["category", "분류", "VARCHAR(50)", "", "", "YES", "", "공연/전시/축제 등"],
                [
                    "place_name",
                    "장소명",
                    "TEXT",
                    "",
                    "",
                    "YES",
                    "",
                    "비정규화. 행사 장소가 places에 없을 수 있음 (외부 공연장)",
                ],
                ["address", "주소", "TEXT", "", "", "YES", "", "비정규화"],
                ["district", "자치구", "VARCHAR(50)", "", "", "YES", "", "비정규화"],
                ["geom", "좌표", "geometry(Point,4326)", "", "", "YES", "", "PostGIS WGS84. 지도 마커 표시용"],
                ["date_start", "시작일", "DATE", "", "", "YES", "", "행사 시작일. 날짜 범위 필터 조건"],
                ["date_end", "종료일", "DATE", "", "", "YES", "", "행사 종료일. 진행 중 판별 기준"],
                ["price", "가격정보", "TEXT", "", "", "YES", "", "무료/유료/금액"],
                ["poster_url", "포스터URL", "TEXT", "", "", "YES", "", "이미지"],
                ["detail_url", "상세URL", "TEXT", "", "", "YES", "", "외부 페이지"],
                ["summary", "요약", "TEXT", "", "", "YES", "", "행사 내용 요약. 벡터 임베딩 대상"],
                ["source", "데이터출처", "VARCHAR(50)", "", "", "YES", "", "8종 source 태그 (culture_event 등)"],
                ["raw_data", "원본데이터", "JSONB", "", "", "YES", "", "원본 API/CSV 전체 필드 보존"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "YES", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "ETL 재적재 시 갱신"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "소프트 삭제"],
            ],
        },
        {
            "num": "2.3",
            "name_ko": "행정동",
            "name_en": "administrative_districts",
            "rows": "427",
            "desc": "서울시 427개 행정동의 마스터 테이블. PostGIS MultiPolygon geom 컬럼으로 ST_Contains 공간 쿼리에 사용.",
            "features": "혼잡도 분석(CROWDEDNESS) — 장소의 좌표로 행정동을 특정하고, 해당 행정동의 생활인구 데이터를 조회",
            "how": "'홍대 지금 붐비나?' → 홍대 좌표를 ST_Contains(admin.geom, point)로 행정동 매칭 → population_stats에서 현재 시간대 인구 조회. "
            "데이터 소스: vuski/admdongkor ver20260201 (GeoJSON, MIT 라이선스).",
            "columns": [
                ["adm_dong_code", "행정동코드", "VARCHAR(20)", "PK (자연키)", "", "NO", "", "8자리 행정동 코드"],
                ["adm_dong_name", "행정동명", "VARCHAR(50)", "", "", "NO", "", "행정동 이름 (연남동, 서교동 등)"],
                ["district", "자치구", "VARCHAR(50)", "", "", "NO", "", "25 서울 자치구"],
                ["geom", "경계", "geometry(MultiPolygon,4326)", "", "", "YES", "", "PostGIS 행정동 폴리곤"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "경계 데이터 갱신 시 변경"],
            ],
        },
        {
            "num": "2.4",
            "name_ko": "생활인구",
            "name_en": "population_stats",
            "rows": "278,880",
            "desc": "서울 열린데이터 생활인구 통계. 행정동별 시간대별 인구 데이터로, 혼잡도 분석 기능에 사용. append-only.",
            "features": "혼잡도 분석(CROWDEDNESS)",
            "how": "시간대(0~23) × 기준일 × 행정동으로 조회. '홍대 주말 저녁 혼잡도'는 해당 행정동 코드 + 주말 날짜 + 18~21시 time_slot 평균으로 산출. "
            "raw_data JSONB에 연령대별/성별 세분화 인구 32 컬럼 전량 보존 → '20대 여성 유동인구' 같은 상세 쿼리 가능. "
            "append-only이므로 과거 데이터 보존 → 시계열 추세 분석 가능.",
            "columns": [
                ["id", "생활인구ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["base_date", "기준일", "DATE", "", "", "NO", "", "인구 측정 기준 날짜"],
                ["time_slot", "시간대", "SMALLINT", "", "", "NO", "", "0~23"],
                [
                    "adm_dong_code",
                    "행정동코드",
                    "VARCHAR(20)",
                    "",
                    "FK→admin_districts",
                    "NO",
                    "",
                    "ON DELETE RESTRICT",
                ],
                ["total_pop", "총인구수", "INT", "", "", "NO", "0", "해당 시간대 총 생활인구 수"],
                ["raw_data", "원본데이터", "JSONB", "", "", "YES", "", "서울시 CSV 32 컬럼 전량"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성. append-only"],
            ],
        },
        {
            "num": "2.5",
            "name_ko": "행정동코드매핑",
            "name_en": "admin_code_aliases",
            "rows": "11",
            "desc": "행안부 행정동 분할/통합 시 구→신 코드 변환 브릿지. 생활인구 CSV(202603)의 구 코드 → admin_districts의 신 코드 매핑.",
            "features": "혼잡도 분석(CROWDEDNESS) — 구 코드 데이터도 신 코드로 변환하여 조회 가능",
            "how": "population_stats에 9개 구 코드(6,048 row)가 skip된 상태. 혼잡도 쿼리 시 JOIN admin_code_aliases ON old_code로 구 데이터도 해석 가능. "
            "11 매핑 row: 동대문·강동 1→2 분할, 강북 6건 재구획, 강남 1건 재코딩.",
            "columns": [
                ["id", "매핑ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["old_code", "구코드", "VARCHAR(20)", "", "", "NO", "", "UNIQUE(old_code, new_code)"],
                ["new_code", "신코드", "VARCHAR(20)", "", "FK→admin_districts", "NO", "", "ON DELETE RESTRICT"],
                ["change_type", "변경유형", "VARCHAR(20)", "", "", "NO", "", "rename/split/merge/new"],
                ["change_note", "변경사유", "TEXT", "", "", "YES", "", "행정동 분할/통합 사유 메모"],
                ["confidence", "신뢰도", "VARCHAR(10)", "", "", "NO", "", "authoritative/high/medium/low"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
            ],
        },
        {
            "num": "2.6",
            "name_ko": "사용자",
            "name_en": "users",
            "rows": "0",
            "desc": "이메일+비밀번호 가입과 Google OAuth 소셜 로그인을 모두 지원하는 사용자 테이블.",
            "features": "회원가입/로그인, Google OAuth, 닉네임/비밀번호 변경, 프로필 수정, 로그아웃",
            "how": "auth_provider='email'이면 password_hash에 bcrypt 해시 필수, google_id NULL. "
            "auth_provider='google'이면 google_id 필수, password_hash NULL. "
            "하나의 테이블에서 두 인증 방식을 관리. "
            "google_id는 Google 계정 고유 식별자(불변), user_oauth_tokens와는 user_id FK로 연결.",
            "columns": [
                ["user_id", "사용자ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["email", "이메일", "VARCHAR(200)", "", "", "NO", "", "UNIQUE"],
                ["password_hash", "비밀번호해시", "VARCHAR(200)", "", "", "YES", "", "bcrypt. email 가입 시 필수"],
                ["auth_provider", "인증방식", "VARCHAR(20)", "", "", "NO", "'email'", "email | google"],
                [
                    "google_id",
                    "구글ID",
                    "VARCHAR(100)",
                    "",
                    "",
                    "YES",
                    "",
                    "Google 계정 고유 식별자. google 가입 시 필수",
                ],
                ["nickname", "닉네임", "VARCHAR(100)", "", "", "YES", "", "UI 표시 이름. 사용자 수정 가능"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "프로필 수정 시 갱신"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "소프트 삭제 (계정 비활성화)"],
            ],
        },
        {
            "num": "2.7",
            "name_ko": "OAuth토큰",
            "name_en": "user_oauth_tokens",
            "rows": "0",
            "desc": "Google OAuth refresh/access token 저장. Calendar API 등 외부 서비스 접근에 사용.",
            "features": "일정 추가(CALENDAR) — Google Calendar API 호출 시 이 테이블의 refresh_token으로 access_token 갱신",
            "how": "scope='calendar.events'인 토큰으로 Google Calendar에 일정 생성. "
            "provider 컬럼은 향후 카카오/네이버 OAuth 확장 대비. "
            "is_deleted로 소프트 삭제하여 토큰 폐기 감사 이력 보존.",
            "columns": [
                ["token_id", "토큰ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["user_id", "사용자ID", "BIGINT", "", "FK→users", "NO", "", "ON DELETE CASCADE"],
                ["provider", "제공자", "VARCHAR(20)", "", "", "NO", "", "google 등. 향후 확장 대비"],
                ["scope", "범위", "VARCHAR(100)", "", "", "NO", "", "OAuth 권한 범위 (calendar.events 등)"],
                [
                    "refresh_token",
                    "리프레시토큰",
                    "VARCHAR(512)",
                    "",
                    "",
                    "NO",
                    "",
                    "장기 토큰. access_token 갱신에 사용",
                ],
                ["access_token", "액세스토큰", "VARCHAR(512)", "", "", "YES", "", "단기 토큰. API 호출 시 사용"],
                ["expires_at", "만료일시", "TIMESTAMPTZ", "", "", "YES", "", "access_token 만료 시각"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "토큰 갱신 시 변경"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "토큰 폐기 이력 보존"],
            ],
        },
        {
            "num": "2.8",
            "name_ko": "대화",
            "name_en": "conversations",
            "rows": "0",
            "desc": "채팅 세션의 메타 정보. 사이드바 목록 UI, 제목 표시에 사용. LangGraph checkpoint와 thread_id로 연동.",
            "features": "채팅 목록/상세 조회, 대화 삭제, 제목 자동생성/수정, 새 대화 시작",
            "how": "thread_id(VARCHAR(100))가 LangGraph checkpoint 및 messages 테이블과의 실질적 연결 키. "
            "LangGraph가 대화 상태를 thread_id로 관리하므로 서비스 전체에서 공통 키로 사용. "
            "conversation_id는 내부 PK(BIGINT), thread_id가 UNIQUE 실제 식별자. "
            "bookmarks/shared_links/feedback에도 thread_id가 있는 것은 의도적 비정규화 (매번 JOIN 회피).",
            "columns": [
                ["conversation_id", "대화ID", "BIGSERIAL", "PK", "", "NO", "auto", "내부 PK"],
                ["thread_id", "스레드ID", "VARCHAR(100)", "", "", "NO", "", "UNIQUE. LangGraph 연동 키"],
                ["user_id", "사용자ID", "BIGINT", "", "FK→users", "NO", "", "ON DELETE CASCADE"],
                ["title", "제목", "VARCHAR(200)", "", "", "YES", "", "LLM 자동 생성 + 사용자 수정 가능"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "제목 수정/메시지 추가 시 갱신"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "소프트 삭제 (대화 삭제)"],
            ],
        },
        {
            "num": "2.9",
            "name_ko": "메시지",
            "name_en": "messages",
            "rows": "0",
            "desc": "대화의 전체 메시지 원본을 영속적으로 저장하는 append-only 테이블. "
            "LLM 컨텍스트 압축과 무관하게 원본 전체를 보존한다. 북마크 위치 이동, 대화 공유, UI 메시지 표시에 사용.",
            "features": "세션 유지, 대화 메시지 조회, 북마크 위치 이동, 대화 공유 조회",
            "how": "이 테이블은 append-only 원칙을 따른다. 한번 INSERT된 레코드는 절대 UPDATE나 DELETE하지 않는다. "
            "따라서 updated_at, is_deleted가 없다. "
            "blocks 컬럼에 16종 WS 콘텐츠 블록(intent/text/text_stream/place/places/events/course/map_markers/"
            "map_route/chart/calendar/references/analysis_sources/disambiguation/done/error)이 JSONB 배열로 저장된다. "
            "LangGraph checkpoint는 LLM 컨텍스트용(압축 가능), messages는 UI 원본(영구 보존) — 이원화 설계.",
            "columns": [
                ["message_id", "메시지ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK. 북마크 위치 참조"],
                [
                    "thread_id",
                    "스레드ID",
                    "VARCHAR(100)",
                    "",
                    "FK→conversations.thread_id",
                    "NO",
                    "",
                    "ON DELETE CASCADE",
                ],
                ["role", "역할", "VARCHAR(20)", "", "", "NO", "", "user | assistant | system"],
                ["blocks", "블록", "JSONB", "", "", "NO", "", "16종 WS 콘텐츠 블록 배열"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성. append-only"],
            ],
        },
        {
            "num": "2.10",
            "name_ko": "북마크",
            "name_en": "bookmarks",
            "rows": "0",
            "desc": "대화 내 특정 위치를 저장하는 북마크. 기존 즐겨찾기(장소 저장) 기능을 대체하여 '대화 위치 저장' 방식으로 재정의. Phase 2.",
            "features": "북마크 생성(5종 프리셋), 전체 목록 조회, 대화별 필터, 위치 이동, 삭제",
            "how": "핀 유형(pin_type)은 place(장소), event(행사), course(코스), analysis(분석), general(일반) 5종. "
            "preview_text에 북마크된 메시지의 미리보기 스니펫을 저장. "
            "북마크 클릭 시 messages 테이블에서 해당 thread_id의 전체 메시지를 불러와서 message_id 위치로 FE 스크롤 이동.",
            "columns": [
                ["bookmark_id", "북마크ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["user_id", "사용자ID", "BIGINT", "", "FK→users", "NO", "", "ON DELETE CASCADE"],
                ["message_id", "메시지ID", "BIGINT", "", "FK→messages", "NO", "", "ON DELETE CASCADE"],
                ["thread_id", "스레드ID", "VARCHAR(100)", "", "", "NO", "", "비정규화 (JOIN 회피)"],
                ["pin_type", "핀유형", "VARCHAR(20)", "", "", "NO", "", "place/event/course/analysis/general"],
                ["preview_text", "미리보기", "TEXT", "", "", "YES", "", "스니펫"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "수정 시 갱신"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "소프트 삭제"],
            ],
        },
        {
            "num": "2.11",
            "name_ko": "공유링크",
            "name_en": "shared_links",
            "rows": "0",
            "desc": "대화 공유 토큰. /shared/{share_token} GET만 인증 우회. Phase 2.",
            "features": "대화 공유 링크 생성, 공유 대화 조회(읽기 전용), 공유 해제",
            "how": "공유 버튼 클릭 시 share_token 생성 → 공유 URL 발급. 전체 대화 공유 또는 from/to_message_id로 범위 지정 가능. "
            "공유받은 사람은 로그인 없이 읽기 전용으로 열람. expires_at으로 만료 관리.",
            "columns": [
                ["share_id", "공유ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["user_id", "사용자ID", "BIGINT", "", "FK→users", "NO", "", "ON DELETE CASCADE"],
                [
                    "from_message_id",
                    "시작메시지ID",
                    "BIGINT",
                    "",
                    "FK→messages",
                    "YES",
                    "",
                    "범위 공유 시. ON DELETE CASCADE",
                ],
                ["to_message_id", "종료메시지ID", "BIGINT", "", "FK→messages", "YES", "", "ON DELETE CASCADE"],
                ["share_token", "공유토큰", "VARCHAR(100)", "", "", "NO", "", "UNIQUE. URL 경로에 노출되는 토큰"],
                ["thread_id", "스레드ID", "VARCHAR(100)", "", "", "NO", "", "비정규화 (JOIN 회피)"],
                ["expires_at", "만료일시", "TIMESTAMPTZ", "", "", "YES", "", "NULL이면 무기한. 만료 관리"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성"],
                ["updated_at", "수정일시", "TIMESTAMPTZ", "", "", "NO", "now()", "공유 해제/갱신 시 변경"],
                ["is_deleted", "삭제여부", "BOOLEAN", "", "", "NO", "false", "소프트 삭제 (공유 해제)"],
            ],
        },
        {
            "num": "2.12",
            "name_ko": "피드백",
            "name_en": "feedback",
            "rows": "0",
            "desc": "AI 응답에 대한 사용자 피드백(👍/👎 + 텍스트). append-only. Phase 3.",
            "features": "AI 응답 피드백",
            "how": "AI 응답 하단의 👍/👎 버튼 클릭 → feedback INSERT. "
            "append-only이므로 updated_at, is_deleted가 없다. 한번 남긴 피드백은 삭제/수정 불가.",
            "columns": [
                ["feedback_id", "피드백ID", "BIGSERIAL", "PK", "", "NO", "auto", "자동 증가 PK"],
                ["user_id", "사용자ID", "BIGINT", "", "FK→users", "NO", "", "ON DELETE CASCADE"],
                ["message_id", "메시지ID", "BIGINT", "", "FK→messages", "NO", "", "ON DELETE CASCADE"],
                ["thread_id", "스레드ID", "VARCHAR(100)", "", "", "NO", "", "비정규화"],
                ["rating", "평가", "VARCHAR(10)", "", "", "NO", "", "up | down"],
                ["comment", "코멘트", "TEXT", "", "", "YES", "", "선택적 텍스트"],
                ["created_at", "생성일시", "TIMESTAMPTZ", "", "", "NO", "now()", "자동 생성. append-only"],
            ],
        },
    ]

    for tbl in TABLES:
        doc.add_page_break()
        doc.add_heading(f"{tbl['num']} {tbl['name_ko']} ({tbl['name_en']}) — {tbl['rows']} row", level=2)

        doc.add_paragraph(tbl["desc"])

        p = doc.add_paragraph()
        p.add_run("연관 기능: ").bold = True
        p.add_run(tbl["features"])

        p2 = doc.add_paragraph()
        p2.add_run("기능 동작 방식: ").bold = True
        p2.add_run(tbl["how"])

        doc.add_paragraph()
        add_table(
            doc,
            ["컬럼", "한글명", "타입", "PK", "FK", "NULL", "기본값", "설명"],
            tbl["columns"],
        )

    # ─── 3. FK 관계 ───
    doc.add_page_break()
    doc.add_heading("3. FK 관계 상세 (12개)", level=1)

    doc.add_paragraph(
        "ON DELETE RESTRICT: 참조 대상 삭제 시 거부 (마스터 데이터 보호). "
        "ON DELETE CASCADE: 부모 삭제 시 자식도 함께 삭제 (사용자→대화→메시지 체인)."
    )

    add_table(
        doc,
        ["#", "FROM 테이블", "FROM 컬럼", "TO 테이블", "TO 컬럼", "ON DELETE"],
        [
            ["1", "population_stats", "adm_dong_code", "administrative_districts", "adm_dong_code", "RESTRICT"],
            ["2", "admin_code_aliases", "new_code", "administrative_districts", "adm_dong_code", "RESTRICT"],
            ["3", "user_oauth_tokens", "user_id", "users", "user_id", "CASCADE"],
            ["4", "conversations", "user_id", "users", "user_id", "CASCADE"],
            ["5", "messages", "thread_id", "conversations", "thread_id", "CASCADE"],
            ["6", "bookmarks", "user_id", "users", "user_id", "CASCADE"],
            ["7", "bookmarks", "message_id", "messages", "message_id", "CASCADE"],
            ["8", "shared_links", "user_id", "users", "user_id", "CASCADE"],
            ["9", "shared_links", "from_message_id", "messages", "message_id", "CASCADE"],
            ["10", "shared_links", "to_message_id", "messages", "message_id", "CASCADE"],
            ["11", "feedback", "user_id", "users", "user_id", "CASCADE"],
            ["12", "feedback", "message_id", "messages", "message_id", "CASCADE"],
        ],
    )

    # ─── 4. OpenSearch ───
    doc.add_page_break()
    doc.add_heading("4. OpenSearch 인덱스 (3개)", level=1)
    doc.add_paragraph("공통: 768d (Gemini gemini-embedding-001), nori 한국어 분석기, k-NN HNSW cosinesimil")

    add_table(
        doc,
        ["인덱스", "PG 연동", "docs", "용도", "임베딩 대상"],
        [
            [
                "places_vector",
                "places.place_id = _id",
                "535,331",
                "장소 의미 검색 (3-Layer page_content)",
                "구조적 속성 + 카테고리 기본 설명 + 리뷰 요약",
            ],
            ["events_vector", "events.event_id = _id", "7,301", "행사 의미 검색", "title + summary"],
            [
                "place_reviews",
                "런타임 생성",
                "~7K (진행 중)",
                "리뷰 키워드 의미 검색",
                "Naver Blog 리뷰 요약 + 6 지표 + 키워드",
            ],
        ],
    )

    # ─── 5. 정규화 ───
    doc.add_heading("5. 정규화 판단", level=1)
    add_table(
        doc,
        ["테이블", "비정규화 컬럼", "사유"],
        [
            ["places", "district", "535K row에서 매번 공간 JOIN 비효율. 자치구 필터링이 가장 빈번한 쿼리 조건"],
            ["events", "district, place_name, address", "행사 장소가 places에 없을 수 있음 (외부 공연장)"],
            ["places/events", "raw_data (JSONB)", "160종 CSV 스키마가 모두 다름. 정규화하면 테이블 수십 개로 폭발"],
            ["bookmarks/shared_links/feedback", "thread_id", "매번 conversations JOIN 회피. 편의상 비정규화"],
        ],
    )

    # ─── 6. 기능·API 매핑 ───
    doc.add_heading("6. 기능·API 매핑 검증", level=1)
    doc.add_paragraph("기능 47개 + API 19개 × ERD 12 테이블 전수 교차 검증 완료. 누락 0건, 불일치 0건.")

    doc.add_heading("6.1 테이블별 커버 기능 수", level=2)
    add_table(
        doc,
        ["테이블", "커버 기능 수", "대표 기능"],
        [
            ["places", "6", "PLACE_SEARCH, PLACE_RECOMMEND, COURSE_PLAN, BOOKING, 지도마커, COST_ESTIMATE"],
            ["events", "3", "EVENT_SEARCH (정형+비정형), COURSE_PLAN, 지도마커"],
            ["administrative_districts", "1", "CROWDEDNESS"],
            ["population_stats", "1", "CROWDEDNESS"],
            ["users", "3", "회원가입/로그인, Google OAuth, 프로필수정"],
            ["user_oauth_tokens", "1", "CALENDAR (Google Calendar API)"],
            ["conversations", "5", "채팅목록, 상세조회, 삭제, 제목수정, 새대화"],
            ["messages", "4", "세션유지, 메시지조회, 북마크위치이동, 대화공유조회"],
            ["bookmarks", "5", "생성, 목록, 필터, 위치이동, 삭제"],
            ["shared_links", "3", "공유생성, 공유조회, 공유해제"],
            ["feedback", "1", "AI 응답 피드백"],
            ["admin_code_aliases", "0 (내부)", "ETL 보조 (혼잡도 쿼리 시 간접 사용)"],
        ],
    )

    # ─── 7. 부록 ───
    doc.add_heading("7. 부록: LangGraph 자동 테이블", level=1)
    doc.add_paragraph(
        "langgraph-checkpoint-postgres 라이브러리가 PostgresSaver.setup() 호출 시 자동 생성하는 4개 테이블. "
        "우리가 DDL을 작성하지 않으며, 라이브러리가 스키마를 관리한다."
    )

    add_table(
        doc,
        ["테이블", "역할"],
        [
            ["checkpoints", "대화 상태 (LLM 컨텍스트) 저장"],
            ["checkpoint_writes", "쓰기 로그"],
            ["checkpoint_blobs", "바이너리 데이터"],
            ["checkpoint_migrations", "스키마 버전 추적 (9 row)"],
        ],
    )
    doc.add_paragraph(
        "이 4개 테이블은 ERDCloud에 포함하지 않는다 (라이브러리 자동관리). "
        "conversations.thread_id가 이 checkpoint 시스템과 연결되는 키이다."
    )

    # Save
    output = Path(__file__).resolve().parents[2] / "기획" / "ERD_v6.3_최종확정보고서.docx"
    doc.save(str(output))
    print(f"saved: {output}")


if __name__ == "__main__":
    build()
