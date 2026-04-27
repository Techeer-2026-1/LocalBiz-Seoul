"""OpenSearch 벡터 DB 구조 확정 문서 docx.

Usage:
    cd backend && source venv/bin/activate
    python scripts/generate_os_structure.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    shading.append(elm)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "1A5276")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F2F2")


def build():
    doc = Document()
    REPO = Path(__file__).resolve().parents[2]

    # Title
    t = doc.add_heading("LocalBiz Intelligence", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_heading("OpenSearch 벡터 DB 구조 확정 문서", level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "v1.0 | 2026-04-13 | 작성: 이정 (BE/PM)\n"
        "인프라: OpenSearch 2.17 (GCE REDACTED_OS_HOST:9200) | 임베딩: Gemini 768d | 분석기: nori"
    )
    r.font.size = Pt(9)

    # ═══ 1. 개요 ═══
    doc.add_page_break()
    doc.add_heading("1. 개요", level=1)

    doc.add_heading("1.1 OpenSearch의 역할", level=2)
    doc.add_paragraph(
        "PostgreSQL이 정형 데이터(SQL 필터, PostGIS 공간 쿼리)를 담당하고, "
        "OpenSearch는 비정형 의미 검색(벡터 k-NN, 키워드 BM25)을 담당한다. "
        "두 DB는 place_id = _id로 application-level 연결."
    )

    add_table(
        doc,
        ["", "PostgreSQL (Cloud SQL)", "OpenSearch (GCE)"],
        [
            ["역할", "정형 저장 + SQL 필터 + 공간 쿼리", "벡터 의미 검색 + 키워드 매칭"],
            ["데이터", "12 테이블 (535K places 등)", "3 인덱스 (768d 벡터)"],
            ["쿼리 예시", "category='카페' AND district='강남구'", "'카공하기 좋은 조용한 카페' 의미 매칭"],
            ["연결 키", "places.place_id / events.event_id", "_id (동일 값)"],
        ],
    )

    doc.add_heading("1.2 인덱스 요약", level=2)
    add_table(
        doc,
        ["인덱스", "docs", "PG 연동", "용도", "적재 상태"],
        [
            [
                "places_vector",
                "535,431",
                "places.place_id = _id",
                "장소 의미 검색 (3-Layer page_content)",
                "적재 중 (~49%)",
            ],
            ["events_vector", "7,301", "events.event_id = _id", "행사 의미 검색 (title+summary)", "완료"],
            ["place_reviews", "~7,000 (목표)", "place_id app-level", "리뷰 키워드 검색", "적재 중 (~851)"],
        ],
    )

    doc.add_heading("1.3 공통 설정", level=2)
    add_table(
        doc,
        ["항목", "값", "설명"],
        [
            ["임베딩 모델", "Gemini gemini-embedding-001", "Google 무료, 1500 RPM"],
            ["벡터 차원", "768d", "불변식 #7: 768d Gemini 고정, OpenAI 금지"],
            ["k-NN 알고리즘", "HNSW (nmslib)", "Hierarchical Navigable Small World"],
            ["유사도 척도", "cosinesimil", "코사인 유사도 (방향 기반)"],
            ["ef_search", "100", "검색 정확도 파라미터"],
            ["분석기", "nori_analyzer", "한국어 형태소 분석 (nori_tokenizer + lowercase)"],
            ["shard / replica", "1 / 0", "단일 노드, PoC 단계"],
        ],
    )

    # ═══ 2. places_vector ═══
    doc.add_page_break()
    doc.add_heading("2. places_vector — 장소 의미 검색", level=1)

    doc.add_paragraph(
        "places 테이블의 535,431건을 page_content(자연어 텍스트)로 변환 후 "
        "768d 벡터로 임베딩하여 적재한 인덱스. "
        "'카공하기 좋은 카페', '데이트 분위기 레스토랑' 같은 비정형 쿼리를 처리한다."
    )

    doc.add_heading("2.1 매핑 스키마", level=2)
    add_table(
        doc,
        ["필드", "OS 타입", "PG 대응", "용도", "검색 방식"],
        [
            ["_id", "(doc ID)", "places.place_id", "PG↔OS 연결키", "—"],
            ["place_id", "keyword", "places.place_id", "필터용 연결키", "정확 매칭 (term)"],
            ["name", "text (nori)", "places.name", "장소명", "형태소 분석 매칭"],
            ["page_content", "text (nori)", "(생성 텍스트)", "3-Layer 자연어 설명", "BM25 키워드 매칭"],
            ["embedding", "knn_vector 768d", "(Gemini 임베딩)", "page_content의 벡터 표현", "k-NN 코사인 유사도"],
            ["category", "keyword", "places.category", "대분류 필터", "정확 매칭 (pre-filter)"],
            ["sub_category", "keyword", "places.sub_category", "소분류 필터", "정확 매칭"],
            ["district", "keyword", "places.district", "자치구 필터", "정확 매칭"],
            ["source", "keyword", "places.source", "데이터 출처", "정확 매칭"],
            ["lat", "float", "ST_Y(places.geom)", "위도 (지도 표시)", "—"],
            ["lng", "float", "ST_X(places.geom)", "경도", "—"],
        ],
    )

    doc.add_heading("2.2 page_content 3-Layer 구조", level=2)
    doc.add_paragraph(
        "page_content 필드는 벡터 검색 품질의 핵심. 3개 Layer로 구성되며, 모든 장소가 Layer 1+2를 보유한다."
    )

    add_table(
        doc,
        ["Layer", "내용", "커버", "생성"],
        [
            [
                "1. 구조적",
                "상호명 + 카테고리 + 소분류 + 지역 + 주소 + raw_data 속성\n"
                "(면적, 좌석수, 진료시간, 요금 등 source별 특화)",
                "535K\n(100%)",
                "source별 특화 템플릿 20종\n(page_content.py)",
            ],
            [
                "2. 카테고리 기본 설명",
                "'혼밥이나 데이트에 적합. 카공하기 좋은 조용한 분위기의 공간...'",
                "535K\n(100%)",
                "Gemini 1회 생성 → 60종 JSON\n(category_descriptions.json)",
            ],
            [
                "3. 리뷰 요약",
                "'바삭한 치킨, 맥주와 완벽한 조합. 주차 가능.'",
                "~7K",
                "Naver Blog 배치 크롤링 → Gemini 요약",
            ],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("page_content 예시 (소상공인 음식점)", level=3)
    p = doc.add_paragraph()
    p.add_run("Layer 1: ").bold = True
    p.add_run(
        "강동구 암사2동에 위치한 치킨 전문점. 60계치킨암사. 업종: 치킨. 분류: 치킨 전문점. 지점: 선사점. 암사동정웅빌딩 1층.\n"
    )
    p.add_run("Layer 2: ").bold = True
    p.add_run(
        "간단하고 빠르게 식사를 해결할 수 있는 곳입니다. 혼밥하기 좋고, 가성비 있는 메뉴로 부담 없이 즐길 수 있습니다.\n"
    )
    p.add_run("주소: ").bold = True
    p.add_run("서울특별시 강동구 상암로3길 8.")

    doc.add_heading("page_content 예시 (병의원 — raw_data 풍부)", level=3)
    p2 = doc.add_paragraph()
    p2.add_run("Layer 1: ").bold = True
    p2.add_run(
        "금천구에 위치한 의원. 가산기대찬의원. 병원분류: 의원. "
        "삼성서울병원 외래교수 출신 구강외과 전문의. 임플란트, 사랑니 발치, 턱관절 치료. "
        "진료: 월~금 09:00-19:30, 토 09:00-15:00.\n"
    )
    p2.add_run("Layer 2: ").bold = True
    p2.add_run("전문 의료진의 진료와 상담을 받을 수 있는 곳입니다.")

    doc.add_heading("2.3 검색 방식", level=2)
    doc.add_paragraph("places_vector는 3가지 검색 방식을 동시에 지원한다:")

    items = [
        (
            "k-NN 벡터 검색",
            "사용자 쿼리를 Gemini 768d로 임베딩 → embedding 필드와 코사인 유사도 계산. "
            "'카공하기 좋은' 같은 의미적 매칭. ~50ms.",
        ),
        (
            "BM25 키워드 검색",
            "page_content의 nori 토큰과 쿼리 키워드 매칭. '성수동'이라는 정확한 지역명 매칭에 강함. ~10ms.",
        ),
        (
            "카테고리 pre-filter",
            "category/district keyword 필드로 정확 필터링 후 벡터 검색. "
            "Intent Router가 추론한 카테고리로 검색 범위 축소.",
        ),
    ]
    for title_text, desc in items:
        p = doc.add_paragraph()
        p.add_run(f"{title_text}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("가중치 비율: ").bold = True
    p3.add_run("k-NN × 1 + BM25 × 2 (category/district boost 3). 카테고리 pre-filter 적용 시 테스트 정확도 9/9 (100%).")

    # ═══ 3. events_vector ═══
    doc.add_page_break()
    doc.add_heading("3. events_vector — 행사 의미 검색", level=1)

    doc.add_paragraph(
        "events 테이블 7,301건의 title + summary를 결합하여 임베딩한 인덱스. "
        "'아이와 갈 만한 체험', '무료 전시' 같은 비정형 행사 검색을 처리한다."
    )

    doc.add_heading("3.1 매핑 스키마", level=2)
    add_table(
        doc,
        ["필드", "OS 타입", "PG 대응", "용도", "비고"],
        [
            ["_id", "(doc ID)", "events.event_id", "PG↔OS 연결키", ""],
            ["event_id", "keyword", "events.event_id", "필터용", ""],
            ["title", "text (nori)", "events.title", "행사명 검색", ""],
            ["description", "text (nori)", "title + summary 결합", "벡터 원본 + BM25", ""],
            ["embedding", "knn_vector 768d", "(Gemini 임베딩)", "의미 검색", ""],
            ["category", "keyword", "events.category", "행사 분류 필터", "공연/전시/축제 등"],
            ["district", "keyword", "events.district", "자치구 필터", ""],
            ["date_start", "date", "events.date_start", "날짜 범위 필터", "k-NN 전 pre-filtering"],
            ["date_end", "date", "events.date_end", "종료일 필터", ""],
            ["source", "keyword", "events.source", "데이터 출처", ""],
        ],
    )

    doc.add_heading("3.2 날짜 pre-filtering", level=2)
    doc.add_paragraph(
        "'이번 주말 전시' 쿼리 시 Intent Router가 날짜를 파싱하고, "
        "events_vector에서 date_start/date_end range filter를 k-NN 전에 적용한다. "
        "지난 행사가 검색되지 않도록 보장."
    )

    doc.add_heading("3.3 개선 계획: 행사 상세 페이지 크롤링", level=2)
    doc.add_paragraph(
        "현재 description이 사무적('코엑스전시장 B홀 전시/미술')이라 "
        "'가족과 함께할 체험 행사' 같은 경험적 쿼리 매칭이 약함. "
        "개선: detail_url(5,466건 보유) 크롤링 → 대상/비용/프로그램 추출 → description 보강."
    )
    add_table(
        doc,
        ["도메인", "건수", "추출 정보"],
        [
            ["yeyak.seoul.go.kr", "1,291", "대상, 비용, 시설, 프로그램"],
            ["sejongpac.or.kr", "301", "장르, 출연진, 티켓, 관람 연령"],
            ["sfac.or.kr", "208", "참여 대상, 무료/유료"],
            ["sema.seoul.go.kr", "148", "전시 설명, 관람 시간"],
            ["기타", "~3,500", "Gemini HTML 구조화 추출"],
        ],
    )

    # ═══ 4. place_reviews ═══
    doc.add_page_break()
    doc.add_heading("4. place_reviews — 리뷰 키워드 검색", level=1)

    doc.add_paragraph(
        "Naver Blog에서 장소별 리뷰를 크롤링 → Gemini로 6지표 채점 + 키워드 추출 → 임베딩. "
        "places_vector의 page_content만으로 부족한 '경험적 쿼리'를 보강한다."
    )

    doc.add_heading("4.1 매핑 스키마", level=2)
    add_table(
        doc,
        ["필드", "OS 타입", "용도", "비고"],
        [
            ["_id", "(doc ID)", "review_{place_id} 형식", ""],
            ["review_id", "keyword", "리뷰 고유 ID", ""],
            ["place_id", "keyword", "PG places 연결", "app-level (FK 아님)"],
            ["place_name", "text (nori)", "장소명 검색", ""],
            ["summary_text", "text (nori)", "리뷰 요약 + 키워드 결합", "벡터 원본 + BM25"],
            ["embedding", "knn_vector 768d", "의미 검색", ""],
            ["keywords", "keyword (array)", "경험적 키워드 태그", "['혼밥', '가성비', '분위기']"],
            ["stars", "float", "6지표 평균 (1~5)", "정렬용"],
            ["source", "keyword", "수집 소스", "naver_blog_batch"],
            ["category", "keyword", "장소 카테고리", "PG에서 복사"],
            ["district", "keyword", "자치구", ""],
            ["analyzed_at", "date", "분석 일시", ""],
        ],
    )

    doc.add_heading("4.2 6지표 채점", level=2)
    add_table(
        doc,
        ["지표", "영문", "설명", "채점 기준"],
        [
            ["만족도", "satisfaction", "전반적 만족감", "Gemini가 리뷰 텍스트에서 1~5점 추론"],
            ["접근성", "accessibility", "교통/위치 편의", "역 근처, 주차 가능 등"],
            ["청결도", "cleanliness", "위생 상태", "깨끗한, 위생적 등 키워드"],
            ["가성비", "value", "가격 대비 만족", "저렴한, 가성비 등"],
            ["분위기", "atmosphere", "공간 분위기", "조용한, 아늑한, 감성 등"],
            ["전문성", "expertise", "서비스/전문 수준", "친절한, 전문적 등"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("주의: ").bold = True
    p.add_run(
        "6지표 점수는 검색에 직접 사용되지 않음. 검색은 keywords 배열 + summary_text 벡터 매칭. "
        "6지표는 레이더 차트 시각화(REVIEW_COMPARE intent) 시 런타임 Gemini 채점 결과와 함께 사용."
    )

    doc.add_heading("4.3 크롤링 파이프라인", level=2)

    steps = [
        ("1. 장소 선별", "카테고리별 quota 배분 (음식점 2,500 / 카페 1,500 / 관광지 1,500 / ...)"),
        ("2. Naver Blog 검색", "'{상호명} {소분류} 후기' 검색. 상호명 relevance 필터 + 광고 필터 + 지역 필터"),
        ("3. Gemini 분석", "리뷰 텍스트 → 6지표 + 키워드 + 요약. 점수 ≤1.0이면 skip (무관 리뷰 차단)"),
        ("4. 임베딩", "summary_text → Gemini 768d 벡터"),
        ("5. OS 적재", "place_reviews _bulk API"),
    ]
    for title_text, desc in steps:
        p = doc.add_paragraph()
        p.add_run(f"{title_text}: ").bold = True
        p.add_run(desc)

    # ═══ 5. 타입 설명 ═══
    doc.add_page_break()
    doc.add_heading("5. OpenSearch 타입 상세", level=1)

    doc.add_heading("5.1 text vs keyword", level=2)
    add_table(
        doc,
        ["타입", "분석", "검색 방식", "용도", "예시"],
        [
            [
                "text (nori)",
                "형태소 분석 → 토큰 분리",
                "부분 매칭 (BM25)",
                "검색 대상 텍스트",
                "'서울특별시 강남구 카페' → [서울특별시, 강남구, 카페]",
            ],
            [
                "keyword",
                "분석 없음 (원문 그대로)",
                "정확 매칭 (term)",
                "필터링/정렬",
                "'음식점' = '음식점'만. '음식' ≠ '음식점'",
            ],
        ],
    )

    doc.add_heading("5.2 knn_vector", level=2)
    doc.add_paragraph(
        "768개의 float 값으로 구성된 벡터. "
        "Gemini gemini-embedding-001이 텍스트를 768차원 공간의 한 점으로 변환. "
        "의미적으로 유사한 텍스트는 벡터 공간에서 가까이 위치. "
        "cosinesimil로 방향 유사도 계산 (크기 무관, 방향만 비교)."
    )

    doc.add_heading("5.3 date", level=2)
    doc.add_paragraph(
        "ISO 8601 날짜 문자열. events_vector에서 range query로 '이번 주말' 같은 날짜 필터링에 사용. "
        "k-NN 검색 전에 pre-filter로 적용하여 지난 행사 제외."
    )

    # ═══ 6. 검색 흐름 ═══
    doc.add_page_break()
    doc.add_heading("6. 검색 파이프라인과 인덱스 활용", level=1)

    search_img = REPO / "기획" / "검색_파이프라인_다이어그램.png"
    if search_img.exists():
        doc.add_picture(str(search_img), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_table(
        doc,
        ["파이프라인 단계", "사용 인덱스/DB", "역할"],
        [
            ["1. Intent Router", "— (Gemini API)", "intent 분류 + 카테고리 추론 + 쿼리 확장"],
            ["2. SQL 필터", "PostgreSQL places/events", "category + district + PostGIS 정확 필터"],
            ["3. Vector k-NN", "places_vector.embedding", "page_content 의미적 유사도 검색"],
            ["4. BM25 보조", "places_vector.page_content", "키워드 정확 매칭 보조"],
            ["5. place_reviews", "place_reviews.keywords/summary_text", "리뷰 기반 경험적 매칭 보강"],
            ["6. LLM Rerank", "— (Gemini API)", "후보 재정렬 (상식 기반)"],
        ],
    )

    doc.add_heading("6.1 기능별 인덱스 사용 매핑", level=2)
    add_table(
        doc,
        ["기능 (Intent)", "places_vector", "events_vector", "place_reviews"],
        [
            ["PLACE_SEARCH", "O (k-NN + BM25)", "", ""],
            ["PLACE_RECOMMEND", "O", "", "O (키워드 가산)"],
            ["EVENT_SEARCH", "", "O (k-NN + date filter)", ""],
            ["COURSE_PLAN", "O", "", ""],
            ["REVIEW_COMPARE", "", "", "O (키워드 + 6지표)"],
            ["COST_ESTIMATE", "", "", ""],
            ["CROWDEDNESS", "", "", ""],
            ["IMAGE_SEARCH", "O (image_embedding, 향후)", "", ""],
            ["ANALYSIS", "", "", "O"],
            ["BOOKING", "", "", ""],
            ["CALENDAR", "", "", ""],
            ["GENERAL", "", "", ""],
        ],
    )

    # ═══ 7. 데이터 흐름 ═══
    doc.add_page_break()
    doc.add_heading("7. 데이터 적재 흐름", level=1)

    sync_img = REPO / "기획" / "OS_PG_동기화_다이어그램.png"
    if sync_img.exists():
        doc.add_picture(str(sync_img), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_table(
        doc,
        ["파이프라인", "입력", "처리", "출력", "소요"],
        [
            [
                "places_vector",
                "PG places 535K",
                "page_content 생성 → Gemini 768d 배치 임베딩",
                "OS places_vector",
                "~4시간",
            ],
            ["events_vector", "PG events 7.3K", "title+summary 결합 → Gemini 768d", "OS events_vector", "~6분"],
            [
                "place_reviews",
                "Naver Blog API 배치",
                "리뷰 크롤링 → Gemini 분석 → 임베딩",
                "OS place_reviews",
                "~30분 (async)",
            ],
        ],
    )

    doc.add_heading("7.1 재적재 정책", level=2)
    reloads = [
        ("places 변경 시", "places_vector 전체 재적재 필요 (page_content 재생성 + 재임베딩)"),
        ("events 변경 시", "events_vector 전체 재적재 (소량이라 ~6분)"),
        ("리뷰 보강 시", "place_reviews 증분 적재 (기존 유지 + 신규 추가)"),
        ("카테고리 설명 변경 시", "Layer 2 변경 → places_vector 전체 재적재"),
    ]
    for title_text, desc in reloads:
        p = doc.add_paragraph()
        p.add_run(f"{title_text}: ").bold = True
        p.add_run(desc)

    # Save
    output = REPO / "기획" / "OpenSearch_벡터DB_구조확정.docx"
    doc.save(str(output))
    print(f"saved: {output}")


if __name__ == "__main__":
    build()
